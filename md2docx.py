import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 设置标题样式
for i in range(1, 5):
    heading_style = doc.styles[f'Heading {i}']
    heading_font = heading_style.font
    heading_font.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if i == 1:
        heading_font.size = Pt(22)
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(0x1a, 0x47, 0x8a)
    elif i == 2:
        heading_font.size = Pt(16)
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(0x2c, 0x5f, 0xa1)
    elif i == 3:
        heading_font.size = Pt(14)
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(0x3a, 0x73, 0xbf)
    elif i == 4:
        heading_font.size = Pt(12)
        heading_font.bold = True

def add_code_block(doc, code_text):
    """添加代码块"""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        # 灰色背景效果通过底纹模拟
        from docx.oxml import OxmlElement
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'f5f5f5')
        shd.set(qn('w:val'), 'clear')
        run._element.rPr.append(shd)

def add_inline_text(paragraph, text):
    """解析行内文本，处理粗体、行内代码"""
    # 处理行内代码
    parts = re.split(r'(`[^`]+`)', text)
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xc7, 0x25, 0x4e)
            from docx.oxml import OxmlElement
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'fff0f0')
            shd.set(qn('w:val'), 'clear')
            run._element.rPr.append(shd)
        else:
            # 处理粗体 **...**
            bold_parts = re.split(r'(\*\*[^*]+\*\*)', part)
            for bp in bold_parts:
                if bp.startswith('**') and bp.endswith('**'):
                    run = paragraph.add_run(bp[2:-2])
                    run.font.bold = True
                else:
                    paragraph.add_run(bp)

def add_table_from_md(doc, lines, start_idx):
    """解析并添加Markdown表格"""
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1

    if len(table_lines) < 2:
        return i

    # 解析表头
    headers = [cell.strip() for cell in table_lines[0].split('|') if cell.strip()]
    # 跳过分隔行（|---|---|）
    data_start = 1
    if table_lines[1].replace('|', '').replace('-', '').replace(' ', '').strip() == '':
        data_start = 2

    # 数据行
    rows = []
    for line in table_lines[data_start:]:
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)

    if not headers:
        return i

    # 创建Word表格
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'

    # 表头
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ''
        run = cell.paragraphs[0].add_run(header)
        run.font.bold = True
        run.font.size = Pt(10)

    # 数据
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            if c < len(headers):
                cell = table.rows[r + 1].cells[c]
                cell.text = ''
                run = cell.paragraphs[0].add_run(cell_text)
                run.font.size = Pt(10)

    doc.add_paragraph()  # 表后空行
    return i


# 读取Markdown文件
with open(r'D:\java project\mall-swarm-master\面试准备-项目深度解析.md', 'r', encoding='utf-8') as f:
    content = f.read()

lines = content.split('\n')
i = 0
in_code_block = False
code_lines = []
in_list = False
list_buffer = []

while i < len(lines):
    line = lines[i]

    # 处理代码块
    if line.strip().startswith('```'):
        if in_code_block:
            # 结束代码块
            add_code_block(doc, '\n'.join(code_lines))
            doc.add_paragraph()  # 代码块后空行
            code_lines = []
            in_code_block = False
        else:
            in_code_block = True
        i += 1
        continue

    if in_code_block:
        code_lines.append(line)
        i += 1
        continue

    # 处理表格
    if line.strip().startswith('|') and not in_code_block:
        # 检查是否是完整表格（下一行也是 | 或者是分隔行）
        if i + 1 < len(lines) and (lines[i+1].strip().startswith('|')):
            i = add_table_from_md(doc, lines, i)
            continue

    # 处理水平分割线
    if line.strip() == '---':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        from docx.oxml import OxmlElement
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:color'), 'cccccc')
        pBdr.append(bottom)
        pPr.append(pBdr)
        i += 1
        continue

    # 处理空行
    if line.strip() == '':
        i += 1
        continue

    # 处理标题
    heading_match = re.match(r'^(#{1,4})\s+(.+)', line)
    if heading_match:
        level = len(heading_match.group(1))
        text = heading_match.group(2).strip()
        # 清理标题中的粗体标记
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        doc.add_heading(text, level=level)
        i += 1
        continue

    # 处理无序列表
    list_match = re.match(r'^(-|\d+\.)\s+(.+)', line)
    if list_match:
        text = list_match.group(2).strip()
        # 清理粗体标记
        clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        p = doc.add_paragraph(style='List Bullet')
        add_inline_text(p, clean_text)
        i += 1
        continue

    # 处理子列表（缩进的）
    sublist_match = re.match(r'^\s{2,4}(-|\d+\.)\s+(.+)', line)
    if sublist_match:
        text = sublist_match.group(2).strip()
        clean_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        p = doc.add_paragraph(style='List Bullet 2')
        add_inline_text(p, clean_text)
        i += 1
        continue

    # 普通段落
    p = doc.add_paragraph()
    # 清理粗体标记后添加 (这里保留粗体渲染)
    add_inline_text(p, line.strip())
    i += 1

# 设置页面边距
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# 保存文档
output_path = r'D:\java project\mall-swarm-master\面试准备-项目深度解析.docx'
doc.save(output_path)
print(f"Word文档已生成: {output_path}")
